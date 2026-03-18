import requests
import tempfile
import os
from typing import Dict, Any, Optional, List
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
from urllib.parse import urlparse
import logging
from PIL import Image
import io
import base64
from dashscope import MultiModalConversation
import zipfile
from docx.oxml import parse_xml
from docx.oxml.ns import qn

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class MultimodalDocxParser:
    """多模态DOCX文档解析器，支持文本、表格和图片分析"""

    def __init__(self, api_key: str = None, vector_api_key: str = None, vector_endpoint: str = None,
                 collection_name: str = None):
        self.api_key = api_key
        self.session = requests.Session()
        # 设置请求头，模拟浏览器访问
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

        # 初始化向量检索器（如果提供了向量服务配置）
        self.vector_retriever = None
        if vector_api_key and vector_endpoint and collection_name:
            from function.dashvector_retriever import DashVectorRetriever
            self.vector_retriever = DashVectorRetriever(api_key, vector_api_key, vector_endpoint, collection_name)

    def parse_docx_structure(self, file_path: str) -> Dict[str, Any]:
        """
        解析docx文档结构，提取四个部分的内容

        Args:
            file_path: docx文件路径

        Returns:
            包含四个部分内容的字典
        """
        try:
            logger.info(f"开始解析文档结构: {file_path}")

            # 打开docx文档
            doc = Document(file_path)
            print("解析文档时，文档的结构doc=", doc)

            # 初始化四个部分的内容
            document_parts = {
                "basic_info": {
                    "text": "",
                    "tables": [],
                    "images": []
                },
                "attention_dimensions": {
                    "text": "",
                    "tables": [],
                    "images": []
                },
                "attention_indicators": {
                    "text": "",
                    "tables": [],
                    "images": []
                },
                "eye_movement_features": {
                    "text": "",
                    "tables": [],
                    "images": []
                }
            }

            # 提取所有内容
            all_content = self._extract_all_content(doc, file_path)
            print("解析文档时，文档的所有内容all_content=", all_content)
            # 分类内容到四个部分
            document_parts = self._classify_content(all_content)
            print("解析文档时，文档的四个部分document_parts=", document_parts)
            logger.info(f"文档结构解析完成")
            return document_parts

        except Exception as e:
            logger.error(f"解析文档结构时发生错误: {str(e)}")
            return {
                "error": f"解析失败: {str(e)}",
                "basic_info": {"text": "", "tables": [], "images": []},
                "attention_dimensions": {"text": "", "tables": [], "images": []},
                "attention_indicators": {"text": "", "tables": [], "images": []},
                "eye_movement_features": {"text": "", "tables": [], "images": []}
            }

    def _extract_all_content(self, doc: Document, docx_path: str) -> Dict[str, Any]:
        """提取文档中的所有内容，包括表格中的图片"""
        content = {
            "paragraphs": [],
            "tables": [],
            "images": []
        }

        # 提取段落文本
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content["paragraphs"].append({
                    "text": text,
                    "style": paragraph.style.name if paragraph.style else "Normal"
                })

        # 提取表格内容和表格中的图片
        for table in doc.tables:
            table_data = []
            table_images = []

            for row_idx, row in enumerate(table.rows):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    # 提取单元格文本
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)

                    # 提取单元格中的图片
                    cell_images = self._extract_images_from_cell(cell, docx_path,
                                                                 f"table_{len(content['tables'])}_{row_idx}_{col_idx}")
                    if cell_images:
                        table_images.extend(cell_images)

                table_data.append(row_data)

            if table_data:
                content["tables"].append({
                    "data": table_data,
                    "images": table_images
                })

        # 提取段落中的图片
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_images = self._extract_images_from_paragraph(paragraph, docx_path, f"para_{para_idx}")
            if para_images:
                content["images"].extend(para_images)

        return content

    def _extract_images_from_cell(self, cell, docx_path: str, prefix: str) -> List[str]:
        """从表格单元格中提取图片"""
        images = []

        # 遍历单元格中的所有元素
        for element in cell._element.iter():
            if element.tag == qn('w:drawing'):
                # 处理绘图对象中的图片
                image_paths = self._extract_images_from_drawing(element, docx_path, prefix)
                images.extend(image_paths)
            elif element.tag == qn('w:pict'):
                # 处理图片对象
                image_paths = self._extract_images_from_pict(element, docx_path, prefix)
                images.extend(image_paths)

        return images

    def _extract_images_from_paragraph(self, paragraph, docx_path: str, prefix: str) -> List[str]:
        """从段落中提取图片"""
        images = []

        # 遍历段落中的所有元素
        for element in paragraph._element.iter():
            if element.tag == qn('w:drawing'):
                # 处理绘图对象中的图片
                image_paths = self._extract_images_from_drawing(element, docx_path, prefix)
                images.extend(image_paths)
            elif element.tag == qn('w:pict'):
                # 处理图片对象
                image_paths = self._extract_images_from_pict(element, docx_path, prefix)
                images.extend(image_paths)

        return images

    def _extract_images_from_drawing(self, drawing_element, docx_path: str, prefix: str) -> List[str]:
        """从绘图对象中提取图片"""
        images = []

        try:
            # 查找图片引用
            for blip in drawing_element.iter(qn('a:blip')):
                embed_attr = blip.get(qn('r:embed'))
                if embed_attr:
                    # 提取图片文件
                    image_path = self._extract_image_from_docx(docx_path, embed_attr, prefix)
                    if image_path:
                        images.append(f"file://{image_path}")

            # 查找链接图片
            for blip in drawing_element.iter(qn('a:blip')):
                link_attr = blip.get(qn('r:link'))
                if link_attr:
                    # 提取链接图片文件
                    image_path = self._extract_image_from_docx(docx_path, link_attr, prefix)
                    if image_path:
                        images.append(f"file://{image_path}")

        except Exception as e:
            logger.warning(f"提取绘图对象中的图片时出错: {str(e)}")

        return images

    def _extract_images_from_pict(self, pict_element, docx_path: str, prefix: str) -> List[str]:
        """从图片对象中提取图片"""
        images = []

        try:
            # 查找图片引用
            for imagedata in pict_element.iter(qn('v:imagedata')):
                src_attr = imagedata.get('src')
                if src_attr:
                    # 提取图片文件
                    image_path = self._extract_image_from_docx(docx_path, src_attr, prefix)
                    if image_path:
                        images.append(f"file://{image_path}")

        except Exception as e:
            logger.warning(f"提取图片对象中的图片时出错: {str(e)}")

        return images

    def _extract_image_from_docx(self, docx_path: str, image_ref: str, prefix: str) -> Optional[str]:
        """从docx文件中提取图片"""
        try:
            # 创建临时目录用于存储图片
            temp_dir = tempfile.mkdtemp(prefix=f"docx_images_{prefix}_")

            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # 查找图片文件
                image_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]

                for image_file in image_files:
                    # 提取图片到临时目录
                    image_name = os.path.basename(image_file)
                    temp_image_path = os.path.join(temp_dir, f"{prefix}_{image_name}")

                    with zip_file.open(image_file) as source, open(temp_image_path, 'wb') as target:
                        target.write(source.read())

                    logger.info(f"提取图片: {temp_image_path}")
                    return temp_image_path

        except Exception as e:
            logger.error(f"从docx文件中提取图片时出错: {str(e)}")

        return None

    def _classify_content(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """将内容分类到四个部分"""
        document_parts = {
            "basic_info": {
                "text": "",
                "tables": [],
                "images": []
            },
            "attention_dimensions": {
                "text": "",
                "tables": [],
                "images": []
            },
            "attention_indicators": {
                "text": "",
                "tables": [],
                "images": []
            },
            "eye_movement_features": {
                "text": "",
                "tables": [],
                "images": []
            }
        }

        # 分类段落文本
        current_section = "basic_info"

        for para in content["paragraphs"]:
            text = para["text"]

            # 根据关键词判断当前部分
            if any(keyword in text for keyword in ["编号", "性别", "年龄", "日期", "基本信息"]):
                current_section = "basic_info"
            elif any(keyword in text for keyword in ["注意维度指标汇总", "五维度模型", "三维度模型"]):
                current_section = "attention_dimensions"
            elif any(keyword in text for keyword in ["注意维度指标图示", "报告结果", "分析结论"]):
                current_section = "attention_indicators"
            elif any(keyword in text for keyword in ["眼动特征", "注视位置", "瞳孔直径", "眼跳速度"]):
                current_section = "eye_movement_features"

            # 将文本添加到当前部分
            if document_parts[current_section]["text"]:
                document_parts[current_section]["text"] += "\n" + text
            else:
                document_parts[current_section]["text"] = text

        # 分类表格
        for table_info in content["tables"]:
            if not table_info or "data" not in table_info:
                continue

            table_data = table_info["data"]
            table_images = table_info.get("images", [])

            # 根据表格内容判断属于哪个部分
            table_text = " ".join([cell for row in table_data for cell in row if cell])

            if any(keyword in table_text for keyword in ["编号", "性别", "年龄", "日期"]):
                document_parts["basic_info"]["tables"].append(table_info)
            elif any(keyword in table_text for keyword in ["指标", "结果", "参考", "偏差", "注意"]):
                document_parts["attention_dimensions"]["tables"].append(table_info)
            elif any(keyword in table_text for keyword in ["注视", "瞳孔", "眼跳", "热图"]):
                document_parts["eye_movement_features"]["tables"].append(table_info)
            else:
                # 默认添加到注意维度指标部分
                document_parts["attention_dimensions"]["tables"].append(table_info)

        return document_parts

    def analyze_with_multimodal_model(self, document_parts: Dict[str, Any], api_key: str) -> str:
        """
        使用qwen-vl-plus多模态模型分析文档内容

        Args:
            document_parts: 文档的四个部分内容
            api_key: 阿里云API密钥

        Returns:
            AI分析结果
        """
        try:
            logger.info("开始使用多模态模型分析文档")

            # 构建系统提示词
            system_prompt = """你是一个专业的眼动测试分析专家。请根据提供的眼动检查报告内容进行专业的分析。

报告内容分为四个部分：
1. 基本信息：包含编号、性别、年龄、日期等
2. 注意维度指标汇总：包含五维度模型和三维度模型的指标数据
3. 注意维度指标图示与报告结果：包含图表和对应的分析结果
4. 眼动特征：包含注视位置、瞳孔直径、眼跳速度等眼动数据

请从以下方面进行分析：
- 基本信息评估：受试者的基本特征
- 注意维度分析：评估各维度的指标是否正常，识别异常项
- 眼动特征分析：评估眼动数据的正常性
- 综合评估：给出整体认知功能评估
- 干预建议：提供针对性的训练和干预建议

请确保分析专业、准确、简洁，控制在200字以内。"""

            # 构建用户消息
            user_message = self._build_multimodal_message(document_parts)

            # 调用多模态模型
            response = MultiModalConversation.call(
                model='qwen-vl-plus',
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                api_key=api_key
            )

            if response.status_code == 200:
                result = response.output.choices[0].message.content
                logger.info("多模态模型分析完成")
                return result
            else:
                logger.error(f"多模态模型调用失败: {response.message}")
                return f"分析失败: {response.message}"

        except Exception as e:
            logger.error(f"多模态模型分析时发生错误: {str(e)}")
            return f"分析过程中出现错误: {str(e)}"

    def _build_multimodal_message(self, document_parts: Dict[str, Any]) -> List[Dict[str, Any]]:
        """构建多模态消息"""
        message_content = []

        # 添加文本内容
        text_content = "请分析以下眼动检查报告内容：\n\n"

        # 第一部分：基本信息
        if document_parts["basic_info"]["text"]:
            text_content += "【第一部分：基本信息】\n"
            text_content += document_parts["basic_info"]["text"] + "\n\n"

        # 第二部分：注意维度指标汇总
        if document_parts["attention_dimensions"]["text"] or document_parts["attention_dimensions"]["tables"]:
            text_content += "【第二部分：注意维度指标汇总】\n"
            if document_parts["attention_dimensions"]["text"]:
                text_content += document_parts["attention_dimensions"]["text"] + "\n"
            if document_parts["attention_dimensions"]["tables"]:
                text_content += "表格数据：\n"
                for i, table_info in enumerate(document_parts["attention_dimensions"]["tables"]):
                    text_content += f"表格{i + 1}：\n"
                    # 将表格转换为markdown格式
                    markdown_table = self._convert_table_to_markdown(table_info["data"])
                    text_content += markdown_table + "\n"
                    # 添加表格中的图片引用
                    if table_info.get("images"):
                        text_content += f"表格{i + 1}包含图片：\n"
                        for img_idx, img_path in enumerate(table_info["images"]):
                            text_content += f"图片{img_idx + 1}: {img_path}\n"
                    text_content += "\n"

        # 第三部分：注意维度指标图示与报告结果
        if document_parts["attention_indicators"]["text"]:
            text_content += "【第三部分：注意维度指标图示与报告结果】\n"
            text_content += document_parts["attention_indicators"]["text"] + "\n\n"

        # 第四部分：眼动特征
        if document_parts["eye_movement_features"]["text"] or document_parts["eye_movement_features"]["tables"]:
            text_content += "【第四部分：眼动特征】\n"
            if document_parts["eye_movement_features"]["text"]:
                text_content += document_parts["eye_movement_features"]["text"] + "\n"
            if document_parts["eye_movement_features"]["tables"]:
                text_content += "眼动特征表格：\n"
                for i, table_info in enumerate(document_parts["eye_movement_features"]["tables"]):
                    text_content += f"表格{i + 1}：\n"
                    # 将表格转换为markdown格式
                    markdown_table = self._convert_table_to_markdown(table_info["data"])
                    text_content += markdown_table + "\n"
                    # 添加表格中的图片引用
                    if table_info.get("images"):
                        text_content += f"表格{i + 1}包含图片：\n"
                        for img_idx, img_path in enumerate(table_info["images"]):
                            text_content += f"图片{img_idx + 1}: {img_path}\n"
                    text_content += "\n"

        message_content.append({"type": "text", "text": text_content})

        # 添加图片内容（如果有的话）
        for section_name, section_content in document_parts.items():
            # 添加段落中的图片
            for image in section_content["images"]:
                if image and image.startswith("file://"):
                    message_content.append({
                        "type": "image",
                        "image": image
                    })

            # 添加表格中的图片
            for table_info in section_content["tables"]:
                if table_info.get("images"):
                    for image in table_info["images"]:
                        if image and image.startswith("file://"):
                            message_content.append({
                                "type": "image",
                                "image": image
                            })

        return message_content

    def _convert_table_to_markdown(self, table_data: List[List[str]]) -> str:
        """将表格数据转换为markdown格式"""
        if not table_data:
            return ""

        markdown_lines = []

        # 添加表头
        header = "| " + " | ".join(table_data[0]) + " |"
        markdown_lines.append(header)

        # 添加分隔线
        separator = "| " + " | ".join(["---"] * len(table_data[0])) + " |"
        markdown_lines.append(separator)

        # 添加数据行
        for row in table_data[1:]:
            # 确保每行的列数与表头一致
            while len(row) < len(table_data[0]):
                row.append("")
            row_data = row[:len(table_data[0])]  # 截取到表头长度
            markdown_row = "| " + " | ".join(row_data) + " |"
            markdown_lines.append(markdown_row)

        return "\n".join(markdown_lines)


    def cleanup_temp_images(self, document_parts: Dict[str, Any]):
        """清理临时图片文件"""
        try:
            cleaned_files = []

            for section_name, section_content in document_parts.items():
                # 清理段落中的图片
                for image_path in section_content["images"]:
                    if image_path.startswith("file://"):
                        actual_path = image_path[7:]  # 移除 "file://" 前缀
                        if os.path.exists(actual_path):
                            try:
                                os.unlink(actual_path)
                                cleaned_files.append(actual_path)
                                logger.info(f"临时图片已清理: {actual_path}")
                            except Exception as e:
                                logger.warning(f"删除图片文件失败 {actual_path}: {str(e)}")

                # 清理表格中的图片
                for table_info in section_content["tables"]:
                    if table_info.get("images"):
                        for image_path in table_info["images"]:
                            if image_path.startswith("file://"):
                                actual_path = image_path[7:]  # 移除 "file://" 前缀
                                if os.path.exists(actual_path):
                                    try:
                                        os.unlink(actual_path)
                                        cleaned_files.append(actual_path)
                                        logger.info(f"临时图片已清理: {actual_path}")
                                    except Exception as e:
                                        logger.warning(f"删除图片文件失败 {actual_path}: {str(e)}")

            # 清理临时目录（如果为空）
            for cleaned_file in cleaned_files:
                temp_dir = os.path.dirname(cleaned_file)
                if temp_dir and os.path.exists(temp_dir):
                    try:
                        # 检查目录是否为空
                        if not os.listdir(temp_dir):
                            os.rmdir(temp_dir)
                            logger.info(f"空临时目录已删除: {temp_dir}")
                    except Exception as e:
                        logger.warning(f"删除临时目录失败 {temp_dir}: {str(e)}")

            logger.info(f"总共清理了 {len(cleaned_files)} 个临时图片文件")

        except Exception as e:
            logger.error(f"清理临时图片时发生错误: {str(e)}")

    def process_eye_file(self, file_path: str, api_key: str) -> Dict[str, Any]:
        """
        处理眼动报告文件的完整流程
        Args:
            file_path: 眼动报告文件路径
            api_key: 阿里云API密钥
        Returns:
            处理结果字典
        """
        document_parts = None
        try:
            logger.info(f"开始处理眼动报告文件: {file_path}")

            # 检查文件是否存在
            if not os.path.exists(file_path):
                return {"error": "文件不存在"}

            # 检查文件扩展名
            if not file_path.lower().endswith('.docx'):
                return {"error": "只支持.docx格式的文件"}

            # 解析文档结构
            document_parts = self.parse_docx_structure(file_path)
            if "error" in document_parts:
                return document_parts

            # 使用多模态模型分析
            analysis_result = self.analyze_with_multimodal_model(document_parts, api_key)
            print("使用多模态模型分析眼动报告时，分析结果analysis_result=", analysis_result)

            return {
                "success": True,
                "document_parts": document_parts,
                "analysis_result": analysis_result
            }

        except Exception as e:
            logger.error(f"处理眼动报告文件时发生错误: {str(e)}")
            return {"error": f"处理失败: {str(e)}"}

        finally:
            # 清理临时图片文件
            if document_parts:
                try:
                    self.cleanup_temp_images(document_parts)
                    logger.info("临时图片文件清理完成")
                except Exception as cleanup_error:
                    logger.error(f"清理临时图片文件时出错: {str(cleanup_error)}")
